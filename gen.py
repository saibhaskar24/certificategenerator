from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import pandas as pd
import os
import win32com.client


df = pd.read_excel('fi.xlsx')
l = []
for col in df.columns:
    series = df[col]
    x=[]
    for i in series:
        x.append(str(i))
    l.append(x)

for ll in range(28):
    for rrr in range(2,6):
        name = l[rrr][ll]
        document = Document('MSP.docx')
        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle',2)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(42)
        obj_font.name = 'Times New Roman'

        for paragraph in document.paragraphs:
            if 'sai' in paragraph.text:
                print(paragraph.text)
                paragraph.text = ''
                paragraph.add_run(name, style = 'CommentsStyle').bold = True
                break
        if(name != "nan"):
            name = name
            try:
                os.mkdir("E:\msp\p\\"+ str(ll))
            except OSError:
                print ("")
            filepath=r'E:\msp\p\\'+ str(ll) + "\\" + name

            document.save(filepath + '.docx')
          
